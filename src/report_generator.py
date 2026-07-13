"""
Report Generation Tools

Professional report generation system for machinery diagnostics.
Reports are saved in the reports/ directory as:
  - **HTML** (default) — Interactive Plotly charts, standalone files.
  - **DOCX** (optional) — Structured Word documents for stakeholders.
    Requires ``python-docx``: ``pip install predictive-maintenance-mcp[docx]``
"""

import itertools
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Any, Optional, List
import json

import numpy as np
from scipy.signal import find_peaks

# Import HTML templates
from .html_templates import (
    create_fft_report,
    create_envelope_report,
    create_iso_report
)

from .config import REPORTS_DIR
from .path_safety import safe_resolve

logger = logging.getLogger(__name__)

# Optional DOCX support
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


#: Process-local sequence for report filenames — the Windows clock can
#: return identical microsecond timestamps for back-to-back calls, so a
#: monotonic counter guarantees uniqueness within the process.
_report_sequence = itertools.count()


def timestamped_report_name(prefix: str, label: str, ext: str = "html") -> str:
    """Build a unique, timestamped report filename.

    Two consecutive runs always produce two distinct files (timestamp +
    monotonic sequence), so a re-run never silently overwrites the
    previous report.

    Args:
        prefix: Report family (e.g. 'fft_spectrum').
        label: Signal/model label; path separators are flattened.
        ext: File extension without the dot.

    Returns:
        Filename like 'fft_spectrum_baseline_1_20260712-153205-000042.html'.
    """
    safe_label = Path(label).stem.replace("/", "_").replace("\\", "_")
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    seq = next(_report_sequence)
    return f"{prefix}_{safe_label}_{stamp}-{seq:06d}.{ext}"


def save_fft_report(
    signal_file: str,
    sampling_rate: float,
    frequencies: np.ndarray,
    magnitudes: np.ndarray,
    signal_data: np.ndarray,
    max_freq: Optional[float] = None,
    num_peaks: int = 15,
    rotation_freq: Optional[float] = None
) -> Dict[str, Any]:
    """
    Generate and save professional FFT spectrum report.
    
    Args:
        signal_file: Signal filename
        sampling_rate: Sampling rate in Hz
        frequencies: Frequency array (positive frequencies)
        magnitudes: Magnitude array
        signal_data: Original signal data
        max_freq: Maximum frequency to display (default: Nyquist)
        num_peaks: Number of peaks to detect and label
        rotation_freq: Optional shaft rotation frequency for harmonic labeling
    
    Returns:
        Dictionary with file path, metadata, and summary
    """
    # Apply frequency limit
    if max_freq is None:
        max_freq = sampling_rate / 2.0
    
    mask = frequencies <= max_freq
    freq_display = frequencies[mask]
    mag_display = magnitudes[mask]
    
    # Convert to dB scale (normalized to max)
    max_mag = np.max(mag_display)
    mag_display_db = 20 * np.log10((mag_display + 1e-12) / max_mag)
    
    # Peak detection
    freq_resolution = frequencies[1] - frequencies[0]
    min_peak_distance = max(1, int(10 / freq_resolution))
    
    peak_indices, properties = find_peaks(
        mag_display_db,
        height=-40,  # Within 40 dB of max
        distance=min_peak_distance
    )
    
    # Sort by magnitude and take top N
    peak_mags_db = properties['peak_heights']
    top_peak_idx = np.argsort(peak_mags_db)[::-1][:num_peaks]
    peak_indices = peak_indices[top_peak_idx]
    
    # Build peaks list with harmonic detection
    peaks = []
    for idx in peak_indices:
        freq = float(freq_display[idx])
        mag_db = float(mag_display_db[idx])
        
        # Check if harmonic of rotation frequency
        note = ""
        if rotation_freq and rotation_freq > 0:
            harmonic_order = round(freq / rotation_freq)
            if abs(freq - harmonic_order * rotation_freq) < rotation_freq * 0.1:
                note = f"Harmonic {harmonic_order}× shaft"
        
        peaks.append({
            'frequency': freq,
            'magnitude_db': mag_db,
            'note': note
        })
    
    # Metadata
    metadata = {
        'signal_file': signal_file,
        'sampling_rate': sampling_rate,
        'num_samples': len(signal_data),
        'duration': len(signal_data) / sampling_rate,
        'max_frequency': max_freq,
        'num_peaks': len(peaks),
        'rotation_freq': rotation_freq,
        'report_type': 'fft_spectrum'
    }
    
    # Generate HTML
    html = create_fft_report(
        signal_file=signal_file,
        sampling_rate=sampling_rate,
        frequencies=freq_display.tolist(),
        magnitudes_db=mag_display_db.tolist(),
        peaks=peaks,
        metadata=metadata
    )
    
    # Save HTML file (timestamped: consecutive runs never overwrite)
    output_file = REPORTS_DIR / timestamped_report_name("fft_spectrum", signal_file)
    output_file.write_text(html, encoding='utf-8')
    
    logger.info(f"FFT report saved: {output_file.name}")
    
    return {
        'file_path': str(output_file.absolute()),
        'file_name': output_file.name,
        'file_size_kb': output_file.stat().st_size / 1024,
        'report_type': 'fft_spectrum',
        'num_peaks_detected': len(peaks),
        'peak_frequencies': [p['frequency'] for p in peaks[:5]],  # Top 5 for summary
        'metadata': metadata,
        'message': f"✓ FFT spectrum report saved: {output_file.name} ({output_file.stat().st_size / 1024:.1f} KB)"
    }


def save_envelope_report(
    signal_file: str,
    sampling_rate: float,
    filter_band: tuple,
    filtered_signal: np.ndarray,
    envelope: np.ndarray,
    env_frequencies: np.ndarray,
    env_magnitudes: np.ndarray,
    bearing_freqs: Optional[Dict[str, float]] = None,
    max_freq: float = 500.0,
    num_peaks: int = 15
) -> Dict[str, Any]:
    """
    Generate and save professional envelope analysis report.
    
    Args:
        signal_file: Signal filename
        sampling_rate: Sampling rate in Hz
        filter_band: (low, high) Hz tuple
        filtered_signal: Bandpass filtered signal
        envelope: Envelope signal
        env_frequencies: Envelope spectrum frequencies
        env_magnitudes: Envelope spectrum magnitudes
        bearing_freqs: Optional dict with BPFO, BPFI, BSF, FTF
        max_freq: Max frequency to display in envelope spectrum
        num_peaks: Number of peaks to detect
    
    Returns:
        Dictionary with file path, metadata, and summary
    """
    # Apply frequency limit
    mask = env_frequencies <= max_freq
    env_freq_display = env_frequencies[mask]
    env_mag_display = env_magnitudes[mask]
    
    # Convert to dB scale (normalized to max)
    max_mag = np.max(env_mag_display)
    env_mag_display_db = 20 * np.log10((env_mag_display + 1e-12) / max_mag)
    
    # Peak detection
    freq_resolution = env_frequencies[1] - env_frequencies[0]
    min_peak_distance = max(1, int(5 / freq_resolution))
    
    peak_indices, properties = find_peaks(
        env_mag_display_db,
        height=-40,
        distance=min_peak_distance
    )
    
    # Sort and take top N
    peak_mags_db = properties['peak_heights']
    top_idx = np.argsort(peak_mags_db)[::-1][:num_peaks]
    peak_indices = peak_indices[top_idx]
    
    # Build peaks list with bearing frequency matching
    peaks = []
    for idx in peak_indices:
        freq = float(env_freq_display[idx])
        mag_db = float(env_mag_display_db[idx])
        
        # Check match with bearing frequencies
        match = ""
        if bearing_freqs:
            for name, bf in bearing_freqs.items():
                if bf and abs(freq - bf) < bf * 0.05:  # Within 5%
                    match = f"≈ {name}"
                    break
        
        peaks.append({
            'frequency': freq,
            'magnitude_db': mag_db,
            'match': match
        })
    
    # Time arrays for plotting (downsample for file size)
    downsample_factor = max(1, len(filtered_signal) // 1000)
    time_data = np.linspace(0, len(filtered_signal) / sampling_rate, len(filtered_signal))
    time_display = time_data[::downsample_factor].tolist()
    filtered_display = filtered_signal[::downsample_factor].tolist()
    envelope_display = envelope[::downsample_factor].tolist()
    
    # Metadata
    metadata = {
        'signal_file': signal_file,
        'sampling_rate': sampling_rate,
        'filter_band': filter_band,
        'num_samples': len(filtered_signal),
        'duration': len(filtered_signal) / sampling_rate,
        'max_frequency': max_freq,
        'num_peaks': len(peaks),
        'bearing_frequencies': bearing_freqs,
        'report_type': 'envelope_analysis'
    }
    
    # Generate HTML
    html = create_envelope_report(
        signal_file=signal_file,
        sampling_rate=sampling_rate,
        filter_band=filter_band,
        time_data=time_display,
        filtered_signal=filtered_display,
        envelope=envelope_display,
        env_freq=env_freq_display.tolist(),
        env_mag_db=env_mag_display_db.tolist(),
        peaks=peaks,
        bearing_freqs=bearing_freqs,
        metadata=metadata
    )
    
    # Save HTML file (timestamped: consecutive runs never overwrite)
    output_file = REPORTS_DIR / timestamped_report_name(
        "envelope_analysis", signal_file
    )
    output_file.write_text(html, encoding='utf-8')
    
    logger.info(f"Envelope report saved: {output_file.name}")
    
    # Summary of matches
    matches_found = [p['match'] for p in peaks if p['match']]
    
    return {
        'file_path': str(output_file.absolute()),
        'file_name': output_file.name,
        'file_size_kb': output_file.stat().st_size / 1024,
        'report_type': 'envelope_analysis',
        'num_peaks_detected': len(peaks),
        'peak_frequencies': [p['frequency'] for p in peaks[:5]],
        'bearing_matches': matches_found,
        'metadata': metadata,
        'message': f"✓ Envelope analysis report saved: {output_file.name} ({output_file.stat().st_size / 1024:.1f} KB)"
    }


def save_iso_report(
    signal_file: str,
    iso_result: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Generate and save professional ISO 20816-3 evaluation report.
    
    Args:
        signal_file: Signal filename
        iso_result: ISO evaluation dict (mapped from assess_severity output)
    
    Returns:
        Dictionary with file path, metadata, and summary
    """
    # Metadata
    metadata = {
        'signal_file': signal_file,
        'rms_velocity': iso_result['rms_velocity'],
        'zone': iso_result['zone'],
        'severity_level': iso_result['severity_level'],
        'machine_group': iso_result['machine_group'],
        'support_type': iso_result['support_type'],
        'report_type': 'iso_20816'
    }
    
    # Generate HTML
    html = create_iso_report(
        signal_file=signal_file,
        iso_result=iso_result,
        metadata=metadata
    )
    
    # Save HTML file (timestamped: consecutive runs never overwrite)
    output_file = REPORTS_DIR / timestamped_report_name("iso_20816", signal_file)
    output_file.write_text(html, encoding='utf-8')
    
    logger.info(f"ISO report saved: {output_file.name}")
    
    return {
        'file_path': str(output_file.absolute()),
        'file_name': output_file.name,
        'file_size_kb': output_file.stat().st_size / 1024,
        'report_type': 'iso_20816',
        'zone': iso_result['zone'],
        'severity': iso_result['severity_level'],
        'rms_velocity': iso_result['rms_velocity'],
        'metadata': metadata,
        'message': f"✓ ISO 20816-3 report saved: {output_file.name} - Zone {iso_result['zone']} ({iso_result['severity_level']})"
    }


def read_report_metadata(file_name: str) -> Dict[str, Any]:
    """
    Read metadata from HTML report without loading entire file.

    Args:
        file_name: Report filename in reports/ directory

    Returns:
        Dictionary with the extracted metadata.

    Raises:
        ValueError: If the filename escapes the reports directory, the
            report does not exist, or the report carries no (valid)
            metadata block.
    """
    # Contain the user-supplied filename before touching the filesystem —
    # otherwise ``../../secret`` turns this into a file-existence/size oracle
    # and a conditional-read primitive.
    try:
        file_path = safe_resolve(REPORTS_DIR, file_name)
    except ValueError:
        logger.warning("Rejected out-of-bounds report filename: %r", file_name)
        # Deliberately terse: no directory listing here (oracle stays closed).
        raise ValueError(f"Invalid report filename: {file_name}") from None

    if not file_path.exists():
        available = [f.name for f in REPORTS_DIR.glob("*.html")]
        raise ValueError(
            f"Report not found: {file_name} — available reports: "
            f"{available if available else 'none'}. "
            f"Use list_html_reports() to see them."
        )

    # Read file and extract JSON metadata
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Find metadata JSON block
    start_marker = '<script type="application/json" id="report-metadata">'
    end_marker = '</script>'

    start_idx = content.find(start_marker)
    if start_idx == -1:
        raise ValueError(
            f"Metadata not found in report {file_name} — the file has no "
            f"embedded report-metadata block."
        )

    start_idx += len(start_marker)
    end_idx = content.find(end_marker, start_idx)

    if end_idx == -1:
        raise ValueError(f"Malformed metadata in report {file_name}")

    metadata_json = content[start_idx:end_idx].strip()
    try:
        metadata = json.loads(metadata_json)
    except json.JSONDecodeError as e:
        raise ValueError(f"Malformed metadata in report {file_name}: {e}") from e

    return {
        'file_name': file_name,
        'file_path': str(file_path.absolute()),
        'file_size_kb': file_path.stat().st_size / 1024,
        'metadata': metadata,
        'message': f"Metadata loaded from {file_name}"
    }


def list_reports() -> List[Dict[str, Any]]:
    """
    List all available HTML reports in reports/ directory.
    
    Returns:
        List of dicts with report information
    """
    reports = []

    for html_file in REPORTS_DIR.glob("*.html"):
        # Skip files without a readable metadata block (legitimate for
        # non-report HTML files sitting in the directory).
        try:
            metadata_info = read_report_metadata(html_file.name)
        except ValueError:
            continue

        meta = metadata_info.get('metadata', {})
        reports.append({
            'file_name': html_file.name,
            'file_size_kb': html_file.stat().st_size / 1024,
            'report_type': meta.get('report_type', 'unknown'),
            'signal_file': meta.get('signal_file', 'unknown'),
            'created': html_file.stat().st_mtime
        })
    
    # Sort by creation time (newest first)
    reports.sort(key=lambda x: x['created'], reverse=True)
    
    return reports


# ============================================================================
# DOCX REPORT GENERATION (optional: requires python-docx)
# ============================================================================

def save_diagnostic_report_docx(
    signal_file: str,
    sections: Dict[str, Any],
    title: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generate a structured Word (.docx) diagnostic report.

    ``sections`` is a dict whose keys define what to include.  Recognised
    keys (all optional):

    - ``statistics``  –  dict with RMS, Kurtosis, Crest Factor …
    - ``fft_peaks``   –  list of ``{frequency, magnitude_db, note}``
    - ``envelope_peaks`` – list of ``{frequency, magnitude_db, match}``
    - ``bearing_frequencies`` – dict with BPFO, BPFI, BSF, FTF
    - ``iso``         –  dict mapped from ``assess_severity`` output
    - ``diagnosis``   –  free-text diagnostic summary (str)

    Returns:
        Dictionary with file path and metadata.

    Raises:
        ValueError: If the optional python-docx dependency is not installed.
    """
    if not HAS_DOCX:
        raise ValueError(
            "python-docx is not installed. Install with: "
            "pip install predictive-maintenance-mcp[docx]"
        )

    doc = DocxDocument()

    # -- styles -----------------------------------------------------------
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- title ------------------------------------------------------------
    heading = title or f"Diagnostic Report — {signal_file}"
    doc.add_heading(heading, level=0)
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  •  "
        f"Signal: {signal_file}"
    )

    # -- 1. Statistics table ----------------------------------------------
    stats = sections.get("statistics")
    if stats and isinstance(stats, dict):
        doc.add_heading("Statistical Summary", level=1)
        table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Parameter", "Value"
        for key, val in stats.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = f"{val:.6g}" if isinstance(val, float) else str(val)

    # -- 2. FFT peaks table -----------------------------------------------
    fft_peaks = sections.get("fft_peaks")
    if fft_peaks:
        doc.add_heading("FFT Spectrum — Top Peaks", level=1)
        table = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Frequency (Hz)", "Magnitude (dB)", "Note"
        for p in fft_peaks:
            row = table.add_row().cells
            row[0].text = f"{p['frequency']:.2f}"
            row[1].text = f"{p['magnitude_db']:.1f}"
            row[2].text = p.get("note", "")

    # -- 3. Envelope peaks table ------------------------------------------
    env_peaks = sections.get("envelope_peaks")
    if env_peaks:
        doc.add_heading("Envelope Analysis — Top Peaks", level=1)
        table = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Frequency (Hz)", "Magnitude (dB)", "Bearing Match"
        for p in env_peaks:
            row = table.add_row().cells
            row[0].text = f"{p['frequency']:.2f}"
            row[1].text = f"{p['magnitude_db']:.1f}"
            row[2].text = p.get("match", "")

    # -- 4. Bearing characteristic frequencies ----------------------------
    bf = sections.get("bearing_frequencies")
    if bf and isinstance(bf, dict):
        doc.add_heading("Bearing Characteristic Frequencies", level=1)
        table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Frequency", "Value (Hz)"
        for name, val in bf.items():
            row = table.add_row().cells
            row[0].text = str(name)
            row[1].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)

    # -- 5. ISO 20816 evaluation ------------------------------------------
    iso = sections.get("iso")
    if iso and isinstance(iso, dict):
        doc.add_heading("ISO 20816-3 Evaluation", level=1)
        zone = iso.get("zone", "?")
        severity = iso.get("severity_level", "?")
        rms_v = iso.get("rms_velocity", 0)
        p = doc.add_paragraph()
        p.add_run(f"Zone {zone}").bold = True
        p.add_run(f"  —  {severity}  (RMS velocity: {rms_v:.3f} mm/s)")
        if iso.get("thresholds"):
            doc.add_paragraph(f"Thresholds: {iso['thresholds']}")

    # -- 6. Free-text diagnosis -------------------------------------------
    diagnosis = sections.get("diagnosis")
    if diagnosis:
        doc.add_heading("Diagnostic Summary", level=1)
        doc.add_paragraph(str(diagnosis))

    # -- save (timestamped: consecutive runs never overwrite) --------------
    output_file = REPORTS_DIR / timestamped_report_name(
        "diagnostic", signal_file, ext="docx"
    )
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))

    logger.info("DOCX report saved: %s", output_file.name)

    return {
        "file_path": str(output_file.absolute()),
        "file_name": output_file.name,
        "file_size_kb": output_file.stat().st_size / 1024,
        "report_type": "diagnostic_docx",
        "sections_included": [k for k in sections if sections[k]],
        "message": f"✓ DOCX report saved: {output_file.name} ({output_file.stat().st_size / 1024:.1f} KB)",
    }
